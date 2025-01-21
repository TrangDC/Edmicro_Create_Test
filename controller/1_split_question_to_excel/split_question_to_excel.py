import tkinter as tk
from tkinter import filedialog, messagebox, ttk, scrolledtext
import os
import logging
import threading
from docx import Document
from tkinter import ttk
from tkinter.scrolledtext import ScrolledText
import openpyxl
import zipfile
import shutil
import re
import xml.etree.ElementTree as ET
from docx.oxml import OxmlElement
from docx.table import _Cell, Table
import requests
import base64
import cloudinary
import cloudinary.uploader
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from wcwidth import wcswidth
from PIL import Image, ImageTk
from io import BytesIO
import tempfile
import json
from datetime import datetime
import google.generativeai as genai

# Configuration
cloudinary.config(
    cloud_name="dgbjb4emp",
    api_key="995523778761239",
    api_secret="kO8fKCyTkgKXvkcBZMcqvrncuTk",
    secure=True
)

GOOGLE_API_KEY = "AIzaSyASlUYT5KMxrtMLVLtUpL7mn4MWOtWf29c" # Thay YOUR_GEMINI_API_KEY bằng API key của bạn
genai.configure(api_key=GOOGLE_API_KEY)
generation_config = {
    "temperature": 0.4,
    "top_p": 1,
    "top_k": 32,
    "max_output_tokens": 4096,
}

model = genai.GenerativeModel(model_name="gemini-2.0-flash-exp",
                            generation_config=generation_config,
                            )

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


def describe_image_with_gemini(image_path):
    """Sử dụng Gemini API để mô tả ảnh."""
    try:
        with open(image_path, "rb") as image_file:
            image_bytes = image_file.read()

        image_part = {
            "mime_type": "image/png",
            "data": base64.b64encode(image_bytes).decode()
        }

        prompt_parts = [
            image_part,
            "Hãy mô tả chi tiết và ngắn gọn về hình ảnh này, không thêm bất kỳ yếu tố cảm xúc hoặc chủ quan nào.",
        ]
        
        response = model.generate_content(prompt_parts)
        response.resolve()
        if response and response.text:
            return response.text
        else:
            logger.warning(f"Gemini không trả về mô tả cho ảnh {image_path}")
            return None
    except Exception as e:
        logger.error(f"Lỗi khi mô tả ảnh bằng Gemini: {e}", exc_info=True)
        return None

def table_to_json(table):
    """Convert a Word table to JSON format."""
    table_data = {}
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    
    rows_data = []
    for row in table.rows[1:]:
         row_data = [cell.text.strip() for cell in row.cells]
         rows_data.append(row_data)
    table_data["type"] = "table"
    table_data["title"] = ""
    table_data["header"] = headers
    table_data["rows"] = rows_data

    return json.dumps(table_data, ensure_ascii=False)

def get_image_map_from_relationships(temp_dir):
    """Extract image relationships from document"""
    rels_path = os.path.join(temp_dir, 'word', '_rels', 'document.xml.rels')
    if not os.path.exists(rels_path):
        logger.warning(f"Relationships file not found at {rels_path}")
        return {}

    tree = ET.parse(rels_path)
    root = tree.getroot()

    ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}

    image_map = {}
    for rel in root.findall('.//r:Relationship', ns):
        rid = rel.get('Id')
        target = rel.get('Target')
        if target and ('media/' in target or 'image' in target.lower()):
            image_map[rid] = os.path.basename(target)
            logger.debug(f"Found image relationship: {rid} -> {target}")

    return image_map

def extract_images_from_docx(docx_file, temp_dir, output_dir):
    """Extract images from Word document and save to files."""
    logger.info("Starting image extraction from docx")
    image_dict = {}

    # Get the name of the docx file without extension to use as the directory name
    docx_filename = os.path.splitext(os.path.basename(docx_file))[0]
    output_path = os.path.join(output_dir, docx_filename)
    os.makedirs(output_path, exist_ok=True)

    with zipfile.ZipFile(docx_file, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)

    image_map = get_image_map_from_relationships(temp_dir)
    logger.debug(f"Found image relationships: {image_map}")

    media_dir = os.path.join(temp_dir, 'word', 'media')
    if os.path.exists(media_dir):
        logger.info(f"Processing media directory: {media_dir}")
        for image_file in os.listdir(media_dir):
            if image_file.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                image_path = os.path.join(media_dir, image_file)
                logger.debug(f"Processing image: {image_path}")

                try:
                    # Save image file
                    new_image_name = f"{os.path.splitext(image_file)[0]}.png"
                    new_image_path = os.path.join(output_path, new_image_name)
                    shutil.copy2(image_path, new_image_path)

                    # Create image URL
                    image_url = create_imgur_url(new_image_path)
                    image_dict[image_file] = image_url
                    logger.debug(f"Successfully processed {image_file}")
                except Exception as e:
                    logger.error(f"Error processing image {image_file}: {str(e)}", exc_info=True)

    return image_dict, image_map

def upload_image_to_cloudinary(image_bytes):
    try:
        response = cloudinary.uploader.upload(
            image_bytes,
        )
        if 'secure_url' in response:
            return response['secure_url']
        else:
            return None
    except Exception as e:
        logger.error(f"Error uploading to Cloudinary: {str(e)}", exc_info=True)
        return None
    
def create_imgur_url(image_path):
    """Uploads the image to Cloudinary and returns the https URL"""
    try:
        with open(image_path, "rb") as image_file:
            image_bytes = image_file.read()
        cloudinary_url = upload_image_to_cloudinary(image_bytes)
        return cloudinary_url
    except Exception as e:
        logger.error(f"Error creating Cloudinary URL {str(e)}", exc_info=True)
        return None
    
def find_images_in_paragraph(paragraph, image_map):
    """Extract image references from a paragraph"""
    images = []

    for run in paragraph._element.findall('.//w:drawing',
                                           {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
        blip = run.find('.//a:blip', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        if blip is not None:
            rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rid in image_map:
                logger.debug(f"Found image in paragraph with rid: {rid}")
                images.append(image_map[rid])

    return images

def is_question_start(text):
    """Check if text starts a new question"""
    return bool(re.search(r"^(Câu|Bài|Problem)\s*([IVX\d]+|[a-z]+)\s*[.:]?", text, re.IGNORECASE))

def is_section_header(text):
    """Check if text is a section header"""
    return bool(re.search(r"\((TN|DS|TNN|TL)\)", text))

def get_section_type(text):
    """Extract section type from header"""
    match = re.search(r"\((TN|DS|TNN|TL)\)", text)
    return match.group(1) if match else None

def table_to_markdown(table):
    """Convert a Word table to markdown format with aligned columns and a separator line"""
    markdown_rows = []
    max_col_widths = []

    total_rows = len(table.rows)
    total_cols = len(table.columns)

    # Initialize max column widths
    max_col_widths = [0] * total_cols

    # First pass: get maximum width for each column
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            max_col_widths[j] = max(max_col_widths[j], wcswidth(cell_text))

    # Create header row
    header_cells = []
    for j in range(total_cols):
        header_cells.append(table.rows[0].cells[j].text.strip().ljust(max_col_widths[j]))
    markdown_rows.append("| " + " | ".join(header_cells) + " |")

    # Create separator row
    separator_cells = []
    for width in max_col_widths:
        separator_cells.append("-" * width)
    markdown_rows.append("|" + "|".join(separator_cells) + "|")

    # Create data rows (start from the second row since we used the first one for header)
    for i in range(1, total_rows):
        cells = []
        for j, cell in enumerate(table.rows[i].cells):
            cell_text = cell.text.strip()
            cells.append(cell_text.ljust(max_col_widths[j]))

        markdown_rows.append("| " + " | ".join(cells) + " |")

    return  "\n" + "[table]\n" + "\n".join(markdown_rows) + "\n[/table]" + "\n"

def get_paragraph_text(paragraph):
    """Get text from paragraph while preserving line breaks and applying custom tags for bold/italic/underline"""
    text_parts = []
    for run in paragraph.runs:
        text = run.text
        # if run.bold:
        #     text = f"[bold]{text}[/bold]"
        # if run.italic:
        #     text = f"[italic]{text}[/italic]"
        # if run.underline:
        #     text = f"[underline]{text}[/underline]"
        if run.element.find('.//w:br', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
            text = text.replace('\n', '\n')  # Preserve line breaks
        text_parts.append(text)
    return ''.join(text_parts)

def extract_table_data(table):
    """Extracts data from a Word table as a list of lists."""
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        table_data.append(row_data)
    return table_data

def process_question_content(element, image_map, current_question, temp_dir, table_label=""):
    """Process either a paragraph or table and return content and images"""
    content = ""
    images = []
    
    if isinstance(element, Table):
        table_name = f"Bảng {table_label}" if table_label else ""
        # table_markdown = table_to_markdown(element)
        table_json = table_to_json(element)
        table_data = extract_table_data(element)
        return table_json, table_data, table_name
    else:  # Paragraph
       content = get_paragraph_text(element)
       for run in element._element.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
           blip = run.find('.//a:blip', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
           if blip is not None:
               rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
               if rid in image_map:
                    image_file = image_map[rid]
                    docx_filename = os.path.splitext(os.path.basename(docx_file))[0]
                    image_path = os.path.join("output_images", docx_filename, f"{os.path.splitext(image_file)[0]}.png") # Path to the saved image

                    if os.path.exists(image_path):
                         description = describe_image_with_gemini(image_path)
                         if description:
                             content += f" [Mô tả ảnh: {description} ] "
                         images.append(image_file)
                    else:
                        logger.warning(f"Không tìm thấy file ảnh ở: {image_path}")
       return content, images, ""

def create_sheet_if_not_exists(wb, sheet_name):
    """Create a new sheet if it doesn't exist and set up headers"""
    if sheet_name not in wb.sheetnames:
        ws = wb.create_sheet(sheet_name)
        
        # Setup headers for question sheets
        if sheet_name != "prompt_template":
            #No headers in Question sheets
            return ws
        # Setup header for prompt_template sheet
        else:
             headers = ["Môn học", "Type", "Prompt"]
             for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True)
             #Set data validation for "Type" column
             dv = DataValidation(type="list", formula1='"TN,DS,TNN"')
             ws.add_data_validation(dv)
             for row in range (2, 1000): #add validation to rows from 2 to 1000
                dv.add(ws.cell(row=row, column=2))
             return ws
    return wb[sheet_name]

def create_prompt_template_sheet(wb):
    """Creates and sets up the 'prompt_template' sheet"""
    ws = create_sheet_if_not_exists(wb, "prompt_template")
    ws.column_dimensions[get_column_letter(3)].width = 100  # Set a large width for the "Prompt" column
    for row in ws.rows:
        ws.row_dimensions[row[0].row].height = 40  # Set height for rows in prompt sheet
    return ws

def extract_content(docx_file, subject_var, grade_var):
    """Extract questions and images to Excel, separated by sections."""
    logger.info(f"Starting content extraction from {docx_file}")

    if not os.path.exists(docx_file):
        raise FileNotFoundError(f"Word file not found: {docx_file}")
    
    temp_dir = "temp_media"
    output_dir = "output_images"
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)

    image_dict, image_map = extract_images_from_docx(docx_file, temp_dir, output_dir)
    logger.info(f"Found {len(image_dict)} images and {len(image_map)} image relationships")

    doc = Document(docx_file)
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # Create the prompt_template sheet
    prompt_sheet = create_prompt_template_sheet(wb)

    # Get prompt data from prompt template file
    # prompt_template_path = prompt_template_entry.get() #Get file path from settings
    prompt_template_path = "E:/Edmicro/Tool_xu_ly_de_thi/prompt/prompt_template.xlsx" # Placeholder, replace with actual file path
    # subject = subject_var.get() # Get selected subject
    prompt_data = read_prompt_data_from_sheet(prompt_template_path, subject_var)

    # Write prompt data to prompt_template sheet if data is found
    if prompt_data:
       for row_index, row_data in enumerate(prompt_data, 2):  # Start from row 2
           for col_index, cell_value in enumerate(row_data, 2):
                prompt_sheet.cell(row=row_index, column=col_index, value=cell_value)

   # Write the subject and grade to the first row
    subject_text = f"Môn {subject_var} lớp {grade_var}"
    prompt_sheet.cell(row=2, column=1, value=subject_text)

    current_section = None
    questions = []
    current_question = []
    current_images = []
    # current_table_name = ""  # Initialize current_table_name
    current_table_names = []

    elements = doc.paragraphs + doc.tables
    elements.sort(key=lambda x: x._element.getparent().index(x._element))
    section_counts = {}
    table_sheet = create_sheet_if_not_exists(wb, "tables")
    table_row = 1
    table_count = 1

    for element in elements:
        if isinstance(element, Table):
            table_name = f"Bảng {table_count}"
            table_json, table_data, _ = process_question_content(element, image_map, current_question, temp_dir, str(table_count))
            if table_data:
                table_sheet.cell(row=table_row, column=1, value=table_name)
                table_sheet.cell(row=table_row, column=2, value="\n".join(current_question).strip())
                for row_index, row in enumerate(table_data):
                    for col_index, cell_value in enumerate(row, 1):
                        table_sheet.cell(row=table_row+row_index, column=col_index+2, value=cell_value)
                table_row += len(table_data) + 2
            current_question.append(table_json)
            # current_table_name = table_name  # Store the current table name
            current_table_names.append(table_name)  # Thêm tên bảng vào list
            table_count += 1
        else:  # Paragraph
            text, paragraph_images, _ = process_question_content(element, image_map, current_question, temp_dir)
            text = text.strip()
            text = text.replace("\t", " ").lstrip()

            if is_section_header(text):
                if current_question and questions:
                    questions.append(("\n".join(current_question).strip(), current_images, current_table_names))
                    current_question = []
                    current_images = []
                    # current_table_name = ""  # Reset table name
                    current_table_names = []  # Reset current_table_names

                section_type = get_section_type(text)
                if section_type:
                    if section_type in section_counts:
                        section_counts[section_type] += 1
                        sheet_name = f"{section_type}_{section_counts[section_type]}"
                    else:
                        section_counts[section_type] = 0
                        sheet_name = section_type

                    if current_section and questions:
                        ws = create_sheet_if_not_exists(wb, current_section)
                        write_questions_to_sheet(ws, questions, image_dict)
                        questions = []

                    current_section = sheet_name
                continue

            if is_question_start(text):
                if current_question:
                    questions.append(("\n".join(current_question).strip(), current_images, current_table_names))
                current_question = [text]
                current_images = []
                # current_table_name = ""  # Reset table name for new question
                current_table_names = []  # Reset current_table_names
            else:
                if text:
                    current_question.append(text)
                if paragraph_images:
                    current_images.extend([image_dict[img] for img in paragraph_images if img in image_dict])

    # Add the last question and process last section
    if current_question:
        questions.append(("\n".join(current_question).strip(), current_images, current_table_names))

    if current_section and questions:
        ws = create_sheet_if_not_exists(wb, current_section)
        write_questions_to_sheet(ws, questions, image_dict)
    elif questions:
        ws = create_sheet_if_not_exists(wb, "No Section")
        write_questions_to_sheet(ws, questions, image_dict)

    docx_dir = os.path.dirname(docx_file)
    docx_filename = os.path.splitext(os.path.basename(docx_file))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    excel_filename = f"{docx_filename}_doc_to_excel_export_{timestamp}.xlsx"
    excel_file = os.path.join(docx_dir, excel_filename)
    wb.save(excel_file)
    logger.info(f"Successfully saved questions and images to {excel_file}")
    # log_text.insert(tk.END, f"Successfully saved questions and images to {excel_file}\n")

    shutil.rmtree(temp_dir)
    return excel_file

def write_questions_to_sheet(ws, questions, image_dict):
    """Write questions to the specified worksheet"""
    current_row = 2
    
    headers = ["Question", "Images", "Table", "Correct_Answer", "OpenAI", "Gemini"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)

    for question, images, table_names in questions:
        # Write question and image names
        cell = ws.cell(row=current_row, column=1, value=question)
        cell.alignment = openpyxl.styles.Alignment(wrap_text=True, vertical='top')
        
        if images:
            ws.cell(row=current_row, column=2, value=";".join(images))

        # # Write the table name to table column if it exists
        # if table_name:
        #     ws.cell(row=current_row, column=3, value=table_name)
        # Write all table names
        if table_names:
            ws.cell(row=current_row, column=3, value=";".join(table_names))  # Nối các tên bảng lại

        current_row += 1
    
    # Adjust column widths and row heights
    ws.column_dimensions[get_column_letter(1)].width = 100 # set Question column width
    ws.column_dimensions[get_column_letter(2)].width = 50  # set Image column width
    ws.column_dimensions[get_column_letter(3)].width = 15  # set table column width
    
    # Set row height to accommodate wrapped text
    for row in ws.rows:
        ws.row_dimensions[row[0].row].height = 100  # Adjust this value as needed

def read_prompt_data_from_sheet(prompt_template_path, subject):
    """Reads prompt data from a specific sheet in the prompt template Excel file."""
    try:
        wb = openpyxl.load_workbook(prompt_template_path)
        if subject not in wb.sheetnames:
            logger.warning(f"Sheet '{subject}' not found in prompt template.")
            return []
        ws = wb[subject]
        data = []
        for row in ws.iter_rows(min_row=2, values_only=True):  # Skip header row
            if all(cell is None for cell in row):
                continue
            data.append(row)
        return data
    except Exception as e:
        logger.error(f"Error reading sheet '{subject}' in prompt template: {e}", exc_info=True)
        return []
    
if __name__ == "__main__":
    # Input parameters
    docx_file = "E:\Edmicro\Tool_giải_đề\Đề cần giải\Đề 3 1.docx" # Replace with your actual Word file
    subject = "Toán"       # Replace with your subject
    grade = "10"          # Replace with your grade
    #Call extract content
    try:
        excel_file = extract_content(docx_file, subject, grade)
        print(f"Successfully processed. Excel file saved to {excel_file}")
    except FileNotFoundError as e:
       print(f"Error: {e}")
    except Exception as e:
        print(f"An error occurred: {e}")