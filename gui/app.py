import sys
import os
from PyQt5.QtWidgets import (QApplication, QWidget, QLabel, QLineEdit,
                             QPushButton, QVBoxLayout, QHBoxLayout, QFileDialog,
                             QComboBox, QTextEdit, QProgressBar, QMessageBox,
                             QSizePolicy,QCompleter, QTabWidget)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from docx import Document
from PyQt5.QtWidgets import QComboBox, QLineEdit
from PyQt5.QtGui import QStandardItemModel, QStandardItem, QFont


class CompleterComboBox(QComboBox):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setEditable(True)
        self.lineEdit().installEventFilter(self)  # Install event filter
        self.completer = QCompleter(self)
        self.completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.setCompleter(self.completer)
    
    def eventFilter(self, obj, event):
        if obj == self.lineEdit() and event.type() == event.KeyPress:
            text = self.lineEdit().text()
            if event.key() == Qt.Key_Backspace and not text:
                self.setCurrentIndex(-1)
                return True
            
        return super().eventFilter(obj, event)


    def setModel(self, model):
        super().setModel(model)
        self.completer.setModel(model)

    def setModelColumn(self, column):
      self.completer.setCompletionColumn(column)
      super().setModelColumn(column)



class DocumentGenerator(QThread):
    finished_signal = pyqtSignal(str)  # Signal khi thread hoàn thành
    log_signal = pyqtSignal(str)  # Signal để gửi log message
    progress_signal = pyqtSignal(int)

    def __init__(self, file_path, num_copies):
        super().__init__()
        self.file_path = file_path
        self.num_copies = num_copies
        self.is_running = True

    def run(self):
        try:
            if not self.is_running:
                self.finished_signal.emit("Đã hủy quá trình.")
                return

            self.log_signal.emit("Bắt đầu xử lý...")

            if not self.file_path:
                self.log_signal.emit("Lỗi: Chưa chọn file docx.")
                self.finished_signal.emit("Lỗi: Chưa chọn file docx.")
                return

            if not os.path.exists(self.file_path):
                self.log_signal.emit(f"Lỗi: Không tìm thấy file {self.file_path}.")
                self.finished_signal.emit(f"Lỗi: Không tìm thấy file {self.file_path}.")
                return

            if not self.num_copies or self.num_copies <= 0:
                self.log_signal.emit("Lỗi: Số lượng bản sao phải lớn hơn 0.")
                self.finished_signal.emit("Lỗi: Số lượng bản sao phải lớn hơn 0.")
                return

            doc = Document(self.file_path)
            total_steps = self.num_copies * len(doc.paragraphs)
            current_step = 0

            for i in range(self.num_copies):
                if not self.is_running:
                    self.finished_signal.emit("Đã hủy quá trình.")
                    return

                for paragraph in doc.paragraphs:
                    if not self.is_running:
                        self.finished_signal.emit("Đã hủy quá trình.")
                        return

                    # Do something with paragraph if needed
                    # Example: print paragraph.text

                    current_step += 1
                    progress = int((current_step / total_steps) * 100)
                    self.progress_signal.emit(progress)

                self.log_signal.emit(f"Đã xử lý bản sao thứ {i + 1}")

            self.log_signal.emit("Hoàn thành xử lý file.")
            self.finished_signal.emit("Thành công!")

        except Exception as e:
            self.log_signal.emit(f"Lỗi: {e}")
            self.finished_signal.emit(f"Lỗi: {e}")

    def stop(self):
        self.is_running = False


class MainTabWidget(QTabWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Document Generator")
        # Thay đổi kích thước ở đây
        self.setGeometry(100, 100, 800, 600) # Tăng chiều rộng lên 800, chiều cao lên 600
        self.setMinimumSize(640, 480)
        
        # Tạo tab Tạo đề
        self.create_tab()
        
    def create_tab(self):
        # Tạo widget nội dung cho tab
        self.tab_widget = QWidget()
        
        self.docx_file_path = ""
        self.num_copies = 0
        self.worker_thread = None

        self.setup_ui()
         # Gọi hàm update_subject_list ngay sau khi tạo giao diện
        self.update_subject_list()
        self.update_num_list()

        # Tạo tab với nội dung là widget vừa tạo
        self.addTab(self.tab_widget, "Tạo đề")

    def setup_ui(self):
         # Font chữ lớn hơn
        font = QFont("Arial", 12) # font chữ Arial, kích thước 12

        # Layout chính
        main_layout = QVBoxLayout()
        
         # Chọn file docx
        docx_layout = QHBoxLayout()
        docx_label = QLabel("Chọn file docx:")
        docx_label.setFont(font)
        self.docx_input = QLineEdit()
        self.docx_input.setFont(font)
        docx_button = QPushButton("Chọn file")
        docx_button.setFont(font)
        docx_button.clicked.connect(self.open_file_dialog)

        # Set size policy cho QLineEdit và QPushButton để chúng mở rộng theo layout
        self.docx_input.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        docx_button.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        
        docx_layout.addWidget(docx_label)
        docx_layout.addWidget(self.docx_input)
        docx_layout.addWidget(docx_button)
        main_layout.addLayout(docx_layout)
        
         # Chọn môn thi
        subject_layout = QHBoxLayout()
        subject_label = QLabel("Chọn môn thi:")
        subject_label.setFont(font)
        # Thay QComboBox bằng CompleterComboBox
        self.subject_combo = CompleterComboBox() 
        self.subject_combo.setFont(font)
        # Set size policy cho QComboBox
        self.subject_combo.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        self.subject_combo.setMinimumWidth(100)
        subject_layout.addWidget(subject_label)
        subject_layout.addWidget(self.subject_combo)
        main_layout.addLayout(subject_layout)
       
        # Chọn số lượng bản sao
        num_layout = QHBoxLayout()
        num_label = QLabel("Số lượng bản sao:")
        num_label.setFont(font)
        self.num_combo = CompleterComboBox()
        self.num_combo.setFont(font)
        # Set size policy cho QComboBox
        self.num_combo.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        self.num_combo.setMinimumWidth(100)
        num_layout.addWidget(num_label)
        num_layout.addWidget(self.num_combo)
        main_layout.addLayout(num_layout)
        
        # Loading Bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setFont(font)
        self.progress_bar.setValue(0)
        main_layout.addWidget(self.progress_bar)

        # Log
        log_label = QLabel("Log:")
        log_label.setFont(font)
        self.log_text = QTextEdit()
        self.log_text.setFont(font)
        self.log_text.setReadOnly(True)
        # set size policy cho log text để nó mở rộng khi resize cửa sổ
        self.log_text.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        main_layout.addWidget(log_label)
        main_layout.addWidget(self.log_text)

        # Button start/stop
        button_layout = QHBoxLayout()
        self.start_button = QPushButton("Bắt đầu")
        self.start_button.setFont(font)
        self.start_button.clicked.connect(self.start_process)
        self.stop_button = QPushButton("Hủy")
        self.stop_button.setFont(font)
        self.stop_button.clicked.connect(self.stop_process)
        self.stop_button.setEnabled(False)

        # Set size policy cho các nút start/stop
        self.start_button.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        self.stop_button.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)

        button_layout.addWidget(self.start_button)
        button_layout.addWidget(self.stop_button)
        main_layout.addLayout(button_layout)
        
        self.docx_input.setMinimumHeight(30) # đặt chiều cao tối thiểu cho input là 30
        self.num_combo.setMinimumHeight(30) # đặt chiều cao tối thiểu cho input là 30
        self.subject_combo.setMinimumHeight(30) # đặt chiều cao tối thiểu cho input là 30
        self.log_text.setMinimumHeight(100) # đặt chiều cao tối thiểu cho log là 100
        self.progress_bar.setMinimumHeight(20) # đặt chiều cao tối thiểu cho progress bar là 20

        self.tab_widget.setLayout(main_layout)

    def open_file_dialog(self):
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(self, "Chọn file docx", "", "Word Documents (*.docx)")
        if file_path:
            self.docx_input.setText(file_path)
            self.docx_file_path = file_path
             # Cập nhật môn học
            self.update_subject_list()
        
    def update_subject_list(self):
        subjects = ["", "Toán", "Vật lý", "Hóa học", "Sinh học", "Ngữ văn", "Lịch sử", "Địa lý"]
        model = self.subject_combo.model()
        if model is None:
            model =  QStandardItemModel(self)
            self.subject_combo.setModel(model)
        
        model.clear()
        for item in subjects:
          qItem = QStandardItem(item)
          model.appendRow(qItem)
       
        self.subject_combo.setModelColumn(0)
        self.subject_combo.setCurrentIndex(0)

    def update_num_list(self):
        num_list = [""] + list(str(i) for i in range(1, 11)) # 1 -> 10
        model = self.num_combo.model()
        if model is None:
            model = QStandardItemModel(self)
            self.num_combo.setModel(model)

        model.clear()
        for item in num_list:
          qItem = QStandardItem(item)
          model.appendRow(qItem)

        self.num_combo.setModelColumn(0)
        self.num_combo.setCurrentIndex(0)

    def start_process(self):
        self.log_text.clear()
        num_text = self.num_combo.currentText()
        try:
           self.num_copies = int(num_text)
        except ValueError:
            QMessageBox.warning(self, "Lỗi", "Số lượng bản sao không hợp lệ")
            return
        self.start_button.setEnabled(False)
        self.stop_button.setEnabled(True)
        self.progress_bar.setValue(0)

        if not self.docx_file_path:
            QMessageBox.warning(self, "Lỗi", "Chưa chọn file docx!")
            self.start_button.setEnabled(True)
            self.stop_button.setEnabled(False)
            return

        self.worker_thread = DocumentGenerator(self.docx_file_path, self.num_copies)
        self.worker_thread.finished_signal.connect(self.process_finished)
        self.worker_thread.log_signal.connect(self.update_log)
        self.worker_thread.progress_signal.connect(self.update_progress)
        self.worker_thread.start()

    def stop_process(self):
        if self.worker_thread and self.worker_thread.isRunning():
            self.worker_thread.stop()
            self.log_text.append("Đang dừng...")

    def update_progress(self, value):
        self.progress_bar.setValue(value)

    def update_log(self, log_message):
        self.log_text.append(log_message)

    def process_finished(self, result_message):
        self.start_button.setEnabled(True)
        self.stop_button.setEnabled(False)
        QMessageBox.information(self, "Kết quả", result_message)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    main_window = MainTabWidget()
    main_window.show()
    sys.exit(app.exec_())